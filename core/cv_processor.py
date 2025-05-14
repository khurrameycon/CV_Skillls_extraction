# core/cv_processor.py
import os
import PyPDF2
import docx
import re
from typing import List, Dict, Any, Tuple
import logging
import tempfile
import traceback
from pdfminer.high_level import extract_text as pdfminer_extract_text

class CVProcessor:
    """Handles extraction and preprocessing of CV text from different file formats"""
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file using multiple methods for robustness"""
        text = ""
        error_messages = []
        
        # Try PyPDF2 first
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            
            # If we got reasonable text, return it
            if len(text.strip()) > 100:  # Arbitrary threshold to check if we got meaningful text
                return text
        except Exception as e:
            error_messages.append(f"PyPDF2 extraction failed: {str(e)}")
        
        # If PyPDF2 fails or returns too little text, try pdfminer
        try:
            text = pdfminer_extract_text(file_path)
            if len(text.strip()) > 0:
                return text
        except Exception as e:
            error_messages.append(f"pdfminer extraction failed: {str(e)}")
        
        # Log all errors if both methods failed
        if not text:
            for msg in error_messages:
                self.logger.error(msg)
            raise Exception(f"Failed to extract text from PDF using multiple methods: {', '.join(error_messages)}")
        
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file"""
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            
            return '\n'.join(full_text)
        except Exception as e:
            self.logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from a TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with a different encoding if utf-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                self.logger.error(f"Error extracting text from TXT {file_path} with latin-1 encoding: {str(e)}")
                raise
        except Exception as e:
            self.logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from TXT: {str(e)}")
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from a file based on its extension"""
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            error_msg = f"Unsupported file format: {ext}"
            self.logger.error(error_msg)
            raise ValueError(error_msg)
    
    def preprocess_cv_text(self, text: str) -> str:
        """Clean and preprocess CV text"""
        if not text or len(text.strip()) == 0:
            raise ValueError("Extracted text is empty after preprocessing")
        
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        # Remove any non-printable characters
        text = ''.join(c for c in text if c.isprintable() or c.isspace())
        
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract common sections from a CV
        This is a more robust implementation to better handle different CV formats
        """
        # Define section markers - patterns commonly found in CVs
        section_patterns = {
            "skills": [
                r'\b(?:technical\s+)?skills\b',
                r'\bcompetenc(?:y|ies)\b',
                r'\bproficienc(?:y|ies)\b',
                r'\bqualifications\b',
                r'\bcore\s+abilities\b',
                r'\bkey\s+skills\b',
                r'\bareas\s+of\s+expertise\b'
            ],
            "experience": [
                r'\bwork\s+experience\b', 
                r'\bemployment\s+(?:history|record)\b',
                r'\bprofessional\s+experience\b',
                r'\bcareer\s+history\b',
                r'\bjob\s+history\b',
                r'\bwork\s+history\b',
                r'\bexperience\b'
            ],
            "education": [
                r'\beducation(?:al)?\s+(?:background|history|qualifications)?\b',
                r'\bacademic\s+(?:background|history|qualifications|record)\b',
                r'\bqualifications\b',
                r'\bdegrees?\b',
                r'\beducation\b'
            ]
        }
        
        # Initial section dictionary
        sections = {
            "skills": "",
            "experience": "",
            "education": "",
            "other": ""
        }
        
        # Split text into lines for processing
        lines = text.split('\n')
        current_section = "other"
        section_text = {k: [] for k in sections.keys()}
        
        # Process each line
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:  # Skip empty lines
                continue
                
            # Check if this line is a section header
            new_section_found = False
            for section, patterns in section_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, line.lower()):
                        current_section = section
                        new_section_found = True
                        break
                if new_section_found:
                    break
            
            # If not a section header, add to current section
            if not new_section_found:
                section_text[current_section].append(line)
        
        # Join the section text
        for section in sections:
            sections[section] = '\n'.join(section_text[section])
        
        # If any main section is empty, try to assign content from "other"
        if sections["other"] and (not sections["skills"] or not sections["experience"] or not sections["education"]):
            other_text = sections["other"]
            
            # Try to extract skills
            if not sections["skills"]:
                skills_extracted = self._extract_skills_from_text(other_text)
                if skills_extracted:
                    sections["skills"] = skills_extracted
            
            # Try to extract education (simpler approach than skills)
            if not sections["education"]:
                education_keywords = ['degree', 'university', 'college', 'school', 'bachelor', 'master', 'phd', 'diploma']
                education_lines = []
                for line in other_text.split('\n'):
                    if any(keyword in line.lower() for keyword in education_keywords):
                        education_lines.append(line)
                if education_lines:
                    sections["education"] = '\n'.join(education_lines)
        
        return sections
    
    def _extract_skills_from_text(self, text: str) -> str:
        """Extract skills from text using common skill-related keywords"""
        skill_keywords = [
            'python', 'java', 'javascript', 'sql', 'html', 'css', 'react', 'angular', 'vue', 
            'node', 'express', 'django', 'flask', 'spring', 'hibernate', 'jquery', 'bootstrap',
            'sass', 'less', 'webpack', 'babel', 'git', 'github', 'gitlab', 'bitbucket', 'svn', 
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'circleci', 'travis',
            'agile', 'scrum', 'kanban', 'jira', 'confluence', 'trello', 'slack', 'teams',
            'linux', 'unix', 'windows', 'macos', 'android', 'ios', 'swift', 'kotlin', 'objective-c',
            'c', 'c++', 'c#', '.net', 'asp.net', 'ruby', 'rails', 'php', 'laravel', 'symfony',
            'wordpress', 'drupal', 'joomla', 'magento', 'shopify', 'woocommerce', 'prestashop',
            'seo', 'sem', 'smm', 'google ads', 'facebook ads', 'instagram ads', 'twitter ads',
            'adobe', 'photoshop', 'illustrator', 'indesign', 'xd', 'figma', 'sketch', 'invision',
            'excel', 'word', 'powerpoint', 'access', 'outlook', 'onedrive', 'sharepoint',
            'mongodb', 'mysql', 'postgresql', 'oracle', 'sql server', 'sqlite', 'redis', 'elasticsearch',
            'communication', 'teamwork', 'leadership', 'problem solving', 'critical thinking',
            'time management', 'organization', 'adaptability', 'flexibility', 'creativity',
            'analytical', 'detail-oriented', 'multitasking', 'decision making', 'negotiation',
            'presentation', 'public speaking', 'writing', 'editing', 'proofreading', 'research'
        ]
        
        skill_lines = []
        for line in text.split('\n'):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in skill_keywords):
                skill_lines.append(line)
        
        return '\n'.join(skill_lines) if skill_lines else ""
    
    def process_cv(self, file_path: str) -> Dict[str, Any]:
        """Process a single CV file"""
        try:
            self.logger.info(f"Processing CV: {file_path}")
            
            # Check if file exists
            if not os.path.exists(file_path):
                error_msg = f"File does not exist: {file_path}"
                self.logger.error(error_msg)
                return {"error": error_msg, "filename": os.path.basename(file_path)}
            
            # Extract text
            text = self.extract_text(file_path)
            
            if not text or len(text.strip()) == 0:
                error_msg = f"No text could be extracted from {file_path}"
                self.logger.error(error_msg)
                return {"error": error_msg, "filename": os.path.basename(file_path)}
            
            # Preprocess text
            processed_text = self.preprocess_cv_text(text)
            
            # Extract sections
            sections = self.extract_sections(processed_text)
            
            # Log success
            self.logger.info(f"Successfully processed CV: {file_path}")
            self.logger.debug(f"Extracted sections: {list(sections.keys())}")
            
            return {
                "filename": os.path.basename(file_path),
                "full_text": processed_text,
                "sections": sections
            }
        except Exception as e:
            error_msg = f"Error processing CV {file_path}: {str(e)}"
            self.logger.error(error_msg)
            self.logger.error(traceback.format_exc())
            return {"error": error_msg, "filename": os.path.basename(file_path)}
    
    def process_batch(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Process a batch of CV files"""
        results = []
        for file_path in file_paths:
            result = self.process_cv(file_path)
            results.append(result)
        return results